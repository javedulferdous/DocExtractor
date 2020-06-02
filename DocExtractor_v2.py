from docx import Document
import docx
import subprocess
import os
searchingWord="Comment:"
allComment=[]
file=[]
counter = 0
newpath = r'C:/Users/jafra/OneDrive/Desktop/New folder (3)/test/result/'
searchingWord="Comment:"
endWord= "Time:"

for filename in os.listdir(os.getcwd()):
    if filename.endswith('.docx'):
        file.append(filename)
        counter+=1
print(file)
print("\nNumber of Doc file: ",counter)


for i in file:
    document = Document(i)
    for paragraph in document.paragraphs:
        if searchingWord in paragraph.text:
            if '' in paragraph.text:
                allComment.append(paragraph.text + "\n")  
                output='\n'.join(allComment)

 
if not os.path.exists(newpath):
    os.makedirs(newpath)
_writeDoc = Document()
_writeDoc.add_heading('Result', 0)
_writeDoc.add_paragraph(output)
_writeDoc.save(newpath+"testFile.docx")
